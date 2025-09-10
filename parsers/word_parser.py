"""
Advanced Word Document Parser for .docx files
Extracts text content, structure, images, tables, and metadata
"""

import docx
from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import re
from typing import Dict, List, Optional, Any
import logging
from dataclasses import dataclass, field
from pathlib import Path
import datetime

@dataclass
class WordParagraph:
    """Information about a paragraph"""
    text: str
    style: str
    alignment: str
    font_name: Optional[str] = None
    font_size: Optional[int] = None
    is_bold: bool = False
    is_italic: bool = False

@dataclass 
class WordTable:
    """Information about a table"""
    rows: int
    columns: int
    data: List[List[str]]
    style: Optional[str] = None

@dataclass
class WordSection:
    """Information about a document section"""
    paragraphs: List[WordParagraph]
    tables: List[WordTable]
    page_break_before: bool = False

@dataclass
class WordAnalysis:
    """Complete Word document analysis"""
    filename: str
    file_size: int
    sections: List[WordSection]
    total_paragraphs: int
    total_words: int
    total_characters: int
    page_count: int
    reading_time: float  # in minutes
    document_properties: Dict[str, Any]
    styles_used: List[str]
    images_count: int
    tables_count: int
    headings: List[Dict[str, Any]]
    hyperlinks: List[str]
    comments_count: int
    revisions_count: int

class AdvancedWordParser:
    """Advanced parser for Word documents (.docx format)"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def parse(self, file_path: str) -> WordAnalysis:
        """Parse a Word document and extract comprehensive information"""
        
        file_path = Path(file_path)
        self.logger.info(f"Parsing Word document: {file_path}")
        
        try:
            # Load the document
            doc = docx.Document(file_path)
            
            # Initialize analysis structure
            analysis = WordAnalysis(
                filename=file_path.name,
                file_size=file_path.stat().st_size,
                sections=[],
                total_paragraphs=0,
                total_words=0,
                total_characters=0,
                page_count=1,  # Approximation
                reading_time=0.0,
                document_properties={},
                styles_used=[],
                images_count=0,
                tables_count=0,
                headings=[],
                hyperlinks=[],
                comments_count=0,
                revisions_count=0
            )
            
            # Extract document properties
            analysis.document_properties = self._extract_properties(doc)
            
            # Process document structure
            section = WordSection(paragraphs=[], tables=[])
            full_text = ""
            styles_set = set()
            
            # Iterate through document elements
            for element in doc.element.body:
                if isinstance(element, CT_P):
                    # Process paragraph
                    paragraph = Paragraph(element, doc)
                    if paragraph.text.strip():  # Only non-empty paragraphs
                        para_info = self._extract_paragraph_info(paragraph)
                        section.paragraphs.append(para_info)
                        full_text += paragraph.text + " "
                        
                        if para_info.style:
                            styles_set.add(para_info.style)
                        
                        # Check if it's a heading
                        if para_info.style and 'heading' in para_info.style.lower():
                            analysis.headings.append({
                                'text': para_info.text,
                                'level': self._extract_heading_level(para_info.style),
                                'style': para_info.style
                            })
                
                elif isinstance(element, CT_Tbl):
                    # Process table
                    table = Table(element, doc)
                    table_info = self._extract_table_info(table)
                    section.tables.append(table_info)
                    analysis.tables_count += 1
            
            # Add the section
            analysis.sections.append(section)
            
            # Calculate text statistics
            words = full_text.split()
            analysis.total_words = len(words)
            analysis.total_characters = len(full_text)
            analysis.total_paragraphs = len(section.paragraphs)
            analysis.reading_time = analysis.total_words / 200.0  # Average reading speed
            analysis.styles_used = list(styles_set)
            
            # Extract other elements
            analysis.hyperlinks = self._extract_hyperlinks(doc)
            analysis.images_count = len(doc.inline_shapes)
            
            # Estimate page count (rough approximation)
            analysis.page_count = max(1, analysis.total_words // 250)
            
            return analysis
            
        except Exception as e:
            self.logger.error(f"Failed to parse Word document: {str(e)}")
            raise
    
    def _extract_properties(self, doc: Document) -> Dict[str, Any]:
        """Extract document properties"""
        properties = {}
        
        try:
            core_props = doc.core_properties
            properties.update({
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'comments': core_props.comments or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
                'last_modified_by': core_props.last_modified_by or '',
                'revision': str(core_props.revision) if core_props.revision else '',
                'category': core_props.category or '',
                'language': core_props.language or ''
            })
        except Exception as e:
            self.logger.warning(f"Failed to extract document properties: {str(e)}")
        
        return properties
    
    def _extract_paragraph_info(self, paragraph: Paragraph) -> WordParagraph:
        """Extract detailed information from a paragraph"""
        
        # Get style information
        style_name = paragraph.style.name if paragraph.style else 'Normal'
        
        # Get font information from first run
        font_name = None
        font_size = None
        is_bold = False
        is_italic = False
        
        if paragraph.runs:
            first_run = paragraph.runs[0]
            if first_run.font:
                font_name = first_run.font.name
                font_size = first_run.font.size.pt if first_run.font.size else None
                is_bold = first_run.font.bold or False
                is_italic = first_run.font.italic or False
        
        return WordParagraph(
            text=paragraph.text,
            style=style_name,
            alignment=str(paragraph.alignment) if paragraph.alignment else 'LEFT',
            font_name=font_name,
            font_size=font_size,
            is_bold=is_bold,
            is_italic=is_italic
        )
    
    def _extract_table_info(self, table: Table) -> WordTable:
        """Extract information from a table"""
        
        rows_count = len(table.rows)
        cols_count = len(table.columns) if table.rows else 0
        
        # Extract table data
        data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            data.append(row_data)
        
        return WordTable(
            rows=rows_count,
            columns=cols_count,
            data=data,
            style=table.style.name if table.style else None
        )
    
    def _extract_hyperlinks(self, doc: Document) -> List[str]:
        """Extract hyperlinks from the document"""
        hyperlinks = []
        
        try:
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if hasattr(run.element, 'rPr') and run.element.rPr is not None:
                        # Look for hyperlink elements
                        hyperlink = run.element.rPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink')
                        if hyperlink is not None:
                            hyperlinks.append(run.text)
        except Exception as e:
            self.logger.warning(f"Failed to extract hyperlinks: {str(e)}")
        
        return hyperlinks
    
    def _extract_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name"""
        if not style_name:
            return 0
        
        # Look for numeric patterns in style name
        match = re.search(r'(\d+)', style_name.lower())
        if match:
            return int(match.group(1))
        
        return 1 if 'heading' in style_name.lower() else 0
    
    def extract_text_only(self, file_path: str) -> str:
        """Extract only the text content from the document"""
        try:
            doc = docx.Document(file_path)
            full_text = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            
            return '\n'.join(full_text)
        
        except Exception as e:
            self.logger.error(f"Failed to extract text: {str(e)}")
            return ""
    
    def export_analysis(self, analysis: WordAnalysis, output_path: str) -> None:
        """Export analysis results to JSON"""
        
        # Convert dataclasses to dict for JSON serialization
        def to_dict(obj):
            if hasattr(obj, '__dict__'):
                return {k: to_dict(v) for k, v in obj.__dict__.items()}
            elif isinstance(obj, list):
                return [to_dict(item) for item in obj]
            elif hasattr(obj, 'isoformat'):  # datetime objects
                return obj.isoformat()
            else:
                return obj
        
        analysis_dict = to_dict(analysis)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(analysis_dict, f, indent=2, default=str)
        
        self.logger.info(f"Word analysis exported to {output_path}")

# Example usage
if __name__ == "__main__":
    parser = AdvancedWordParser()
    
    # Test with a Word document
    test_file = "sample.docx"
    
    if Path(test_file).exists():
        try:
            analysis = parser.parse(test_file)
            print(f"\nAnalysis for {test_file}:")
            print(f"Words: {analysis.total_words}")
            print(f"Paragraphs: {analysis.total_paragraphs}")
            print(f"Tables: {analysis.tables_count}")
            print(f"Reading time: {analysis.reading_time:.1f} minutes")
            
            # Export detailed analysis
            parser.export_analysis(analysis, f"{test_file}_analysis.json")
            
        except Exception as e:
            print(f"Error parsing {test_file}: {str(e)}")
    else:
        print(f"Test file {test_file} not found")